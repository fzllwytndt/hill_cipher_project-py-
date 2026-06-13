from docx import Document

def read_docx(file_path):
    """
    Membaca isi file .docx dan menggabungkannya jadi satu string.
    """
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def save_docx(text, file_path):
    """
    Menyimpan string teks ke file .docx baru.
    """
    doc = Document()
    doc.add_paragraph(text)
    doc.save(file_path)

def read_ciphertext(file_path):
    """
    Membaca isi file ciphertext.txt.
    """
    with open(file_path, "r") as f:
        return f.read()

def save_ciphertext(text, file_path):
    """
    Menyimpan teks terenkripsi ke ciphertext.txt.
    """
    with open(file_path, "w") as f:
        f.write(text)
